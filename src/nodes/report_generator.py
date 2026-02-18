"""
MICRA — Report Generator Node
================================

LEARNING CONCEPT: Grounded generation → shareable artifact

This node does something conceptually important: it converts the
accumulated structured state (JSON outputs from 6+ agents) into
a human-readable Word document.

This is the culmination of the "grounded generation" principle:
  Every section of the report is sourced from structured agent outputs.
  Every claim the report makes can be traced back to a source chunk.

This is fundamentally different from asking an LLM to "write a report
about DERMS". That produces fluent but unverifiable text. THIS approach:
  → Facts come from TAM agent output (which cites source chunks)
  → Competitor data comes from competitive intel output (cited)
  → Recommendations come from synthesis output (reasoned from above)

The report generator is NOT a creative writer. It's a formatter.

LEARNING CONCEPT: python-docx for structured document generation

python-docx works like this:
  doc = Document()
  doc.add_heading("Title", level=0)    # Heading 1
  doc.add_paragraph("Body text")
  doc.add_heading("Section", level=1)  # Heading 2
  table = doc.add_table(rows=1, cols=3)
  doc.save("report.docx")

Key formatting objects:
  - Heading levels 0-4 (like h1-h5 in HTML)
  - Paragraph with runs (each run can have different bold/color/size)
  - Table with cells (each cell contains paragraphs)
  - doc.add_page_break() between major sections

LEARNING CONCEPT: Graceful degradation with partial state

Some framework agents may have failed. Some competitor profiles may be empty.
The report generator must handle every combination of present/absent data
gracefully. The pattern: check before rendering, use placeholder text if missing.

A report with "TAM analysis unavailable — framework failed during analysis"
is more useful than a crashed pipeline.
"""

import os
import re
import json
from datetime import datetime
from typing import Any

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from rich.console import Console

from src.state import MICRAState, FrameworkOutput, CompetitorProfile

console = Console()

# ── Color palette ──────────────────────────────────────────────────────────
DARK_BLUE = RGBColor(0x1A, 0x3C, 0x5C)
MEDIUM_BLUE = RGBColor(0x2E, 0x6D, 0xA8)
LIGHT_GRAY = RGBColor(0xF5, 0xF5, 0xF5)
GREEN = RGBColor(0x27, 0x7A, 0x40)
RED = RGBColor(0xC0, 0x39, 0x2B)
ORANGE = RGBColor(0xE6, 0x7E, 0x22)


# ── Helpers ────────────────────────────────────────────────────────────────

def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    """Add a styled heading."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = DARK_BLUE


def _add_body(doc: Document, text: str, bold: bool = False) -> None:
    """Add a body paragraph."""
    p = doc.add_paragraph(text)
    if bold:
        for run in p.runs:
            run.bold = True
    return p


def _add_bullet(doc: Document, text: str, color: RGBColor | None = None) -> None:
    """Add a bullet-point paragraph."""
    p = doc.add_paragraph(text, style="List Bullet")
    if color:
        for run in p.runs:
            run.font.color.rgb = color


def _add_kv(doc: Document, key: str, value: str) -> None:
    """Add a key: value line (key in bold)."""
    p = doc.add_paragraph()
    run_key = p.add_run(f"{key}: ")
    run_key.bold = True
    run_key.font.color.rgb = DARK_BLUE
    p.add_run(value)


def _find_framework(name: str, outputs: list[FrameworkOutput]) -> dict[str, Any] | None:
    """Find a specific framework's output dict, or None."""
    for o in outputs:
        if o["framework_name"] == name:
            return o["output"]
    return None


def _rating_color(rating: str) -> RGBColor:
    """Color-code High/Medium/Low ratings."""
    r = rating.lower()
    if "high" in r:
        return RED
    if "medium" in r:
        return ORANGE
    return GREEN


def _collect_all_sources(state: MICRAState) -> list[dict]:
    """
    Collect all unique sources cited across framework + competitor outputs.

    LEARNING: Citation tracing
    Each FrameworkOutput has source_chunks (list of chunk IDs).
    Each SourceDocument in state["sources"] has chunk_ids + url + title.
    We build a map: chunk_id → source, then deduplicate by URL.

    This is how "every claim cites a source" is enforced:
    frameworks store which chunks they used, we map those back to URLs.
    """
    # Try chunk-level citation tracing first (only works when chunk_ids
    # are populated in SourceDocuments — requires embedder to write them back).
    chunk_to_source: dict[str, dict] = {}
    for source in state.get("sources", []):
        for chunk_id in source.get("chunk_ids", []):
            chunk_to_source[chunk_id] = {
                "url": source["url"],
                "title": source["title"],
                "type": source["source_type"],
            }

    seen_urls: set[str] = set()
    cited_sources: list[dict] = []

    if chunk_to_source:
        # Chunk-level tracing available — list only actually-cited sources
        all_chunk_ids: list[str] = []
        for output in state.get("framework_outputs", []):
            all_chunk_ids.extend(output.get("source_chunks", []))
        for profile in state.get("competitor_profiles", []):
            all_chunk_ids.extend(profile.get("source_chunks", []))

        for chunk_id in all_chunk_ids:
            source = chunk_to_source.get(chunk_id)
            if source and source["url"] not in seen_urls:
                seen_urls.add(source["url"])
                cited_sources.append(source)
    else:
        # Fallback: chunk_ids not populated — list all ingested sources.
        # All sources were scraped for this research run so all are relevant.
        for source in state.get("sources", []):
            if source["url"] not in seen_urls:
                seen_urls.add(source["url"])
                cited_sources.append({
                    "url": source["url"],
                    "title": source["title"],
                    "type": source["source_type"],
                })

    return cited_sources


def _add_labeled_para(doc: Document, label: str, text: str,
                      label_color: RGBColor = None) -> None:
    """Add a paragraph with a colored label run followed by body text."""
    p = doc.add_paragraph()
    r = p.add_run(f"{label}  ")
    r.bold = True
    r.font.color.rgb = label_color or DARK_BLUE
    p.add_run(text)


def _table_header(table, labels: list[str]) -> None:
    """Style the first row of a table as a bold header."""
    hdr = table.rows[0].cells
    for cell, label in zip(hdr, labels):
        cell.text = label
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = DARK_BLUE


# ── Section renderers ──────────────────────────────────────────────────────

def _render_title_page(doc: Document, state: MICRAState) -> None:
    plan = state.get("research_plan", {})
    market = plan.get("target_market", "Market Intelligence")
    geography = plan.get("geography", "")
    date_str = datetime.now().strftime("%B %Y")

    doc.add_paragraph()  # top padding
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(f"{market}")
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = DARK_BLUE

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_para.add_run("Competition Analysis & MVP Strategy")
    sub_run.font.size = Pt(16)
    sub_run.font.color.rgb = MEDIUM_BLUE

    if geography:
        geo_para = doc.add_paragraph(geography)
        geo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in geo_para.runs:
            run.font.size = Pt(13)
            run.font.color.rgb = MEDIUM_BLUE

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Generated by MICRA  •  {date_str}").font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_page_break()


def _render_executive_summary(doc: Document, state: MICRAState) -> None:
    _add_heading(doc, "Executive Summary", level=1)

    bbp = state.get("build_buy_partner_decision", {})
    mvp = state.get("mvp_recommendation", {})
    tam = _find_framework("tam_sam_som", state.get("framework_outputs", []))
    profiles = state.get("competitor_profiles", [])
    plan = state.get("research_plan", {})

    # Strategic questions
    _add_heading(doc, "Key Strategic Questions", level=2)
    sub_questions = plan.get("sub_questions", [])
    if sub_questions:
        for q in sub_questions[:5]:
            _add_bullet(doc, q)
    else:
        for q in [
            "Should we build, buy, or partner to enter this market?",
            "What differentiates us from established players?",
            "What features must the MVP include to win customers?",
            "Which market and geography offers the fastest path to revenue?",
        ]:
            _add_bullet(doc, q)

    # Recommendation
    _add_heading(doc, "Our Recommendation", level=2)
    if bbp.get("recommendation"):
        p = doc.add_paragraph()
        r = p.add_run(f"{bbp['recommendation']}")
        r.font.size = Pt(13)
        r.bold = True
        r.font.color.rgb = DARK_BLUE
        _add_body(doc, bbp.get("reasoning", ""))
        _add_kv(doc, "Time to Market", bbp.get("time_to_market_assessment", ""))

    # Market opportunity snapshot
    _add_heading(doc, "Market Opportunity", level=2)
    if tam:
        _add_kv(doc, "Total Addressable Market", tam.get("tam_value", "—"))
        _add_kv(doc, "Serviceable Market", tam.get("sam_value", "—"))
        _add_kv(doc, "Obtainable Market (3 yr)", tam.get("som_value", "—"))
        _add_kv(doc, "Growth Rate", tam.get("growth_rate", "—"))

    # Competitive landscape snapshot
    _add_heading(doc, "Competitive Landscape", level=2)
    if profiles:
        _add_body(doc, f"{len(profiles)} competitors analyzed: " +
                  ", ".join(p["name"] for p in profiles[:5]))

    # MVP north star
    if mvp.get("north_star_metric"):
        _add_heading(doc, "North Star Metric", level=2)
        p = doc.add_paragraph()
        r = p.add_run(mvp["north_star_metric"])
        r.bold = True
        r.font.color.rgb = MEDIUM_BLUE

    doc.add_page_break()


def _render_market_overview(doc: Document, state: MICRAState) -> None:
    _add_heading(doc, "1. Market Overview & Opportunity", level=1)
    _add_body(doc, state.get("research_brief", ""))

    tam = _find_framework("tam_sam_som", state.get("framework_outputs", []))
    gtm = state.get("gtm_strategy", {})

    # 1.1 Global market size
    if tam:
        _add_heading(doc, "1.1 Global Market Size", level=2)
        table = doc.add_table(rows=4, cols=3)
        table.style = "Table Grid"
        _table_header(table, ["Metric", "Value", "What it means"])
        for i, (metric, val_key, desc_key) in enumerate([
            ("Total Addressable Market", "tam_value", "tam_description"),
            ("Serviceable Addressable Market", "sam_value", "sam_description"),
            ("Serviceable Obtainable Market", "som_value", "som_description"),
        ], start=1):
            row = table.rows[i].cells
            row[0].text = metric
            row[1].text = tam.get(val_key, "—")
            row[2].text = tam.get(desc_key, "")
        doc.add_paragraph()
        _add_kv(doc, "CAGR", tam.get("growth_rate", "—"))

        _add_heading(doc, "1.2 Key Growth Drivers", level=2)
        for trend in tam.get("key_trends_driving_growth", []):
            _add_bullet(doc, trend)

        _add_heading(doc, "1.3 Key Assumptions", level=2)
        for assumption in tam.get("key_assumptions", []):
            _add_bullet(doc, assumption, color=RGBColor(0x55, 0x55, 0x55))

    # 1.4 Regional market analysis (from GTM strategy)
    regions = gtm.get("go_to_market_by_region", [])
    if regions:
        _add_heading(doc, "1.4 Regional Market Analysis", level=2)
        for region_data in regions:
            region = region_data.get("region", "")
            _add_heading(doc, region, level=3)
            chars = region_data.get("market_characteristics", [])
            for c in chars:
                _add_bullet(doc, c)
            if region_data.get("typical_deal_size"):
                _add_kv(doc, "Typical Deal Size", region_data["typical_deal_size"])
            if region_data.get("sales_cycle_months"):
                _add_kv(doc, "Sales Cycle", region_data["sales_cycle_months"])

    doc.add_page_break()


def _render_competitive_landscape(doc: Document, state: MICRAState) -> None:
    _add_heading(doc, "2. Competitive Landscape", level=1)

    # 2.1 Industry dynamics (Porter — no label)
    porter = _find_framework("porter_5_forces", state.get("framework_outputs", []))
    if porter:
        _add_heading(doc, "2.1 Industry Competitive Dynamics", level=2)
        _add_kv(doc, "Overall Market Attractiveness",
                porter.get("overall_market_attractiveness", "—"))

        forces = [
            ("Threat of New Entrants", "threat_of_new_entrants"),
            ("Buyer Power", "bargaining_power_buyers"),
            ("Supplier Power", "bargaining_power_suppliers"),
            ("Threat of Substitutes", "threat_of_substitutes"),
            ("Competitive Rivalry", "competitive_rivalry"),
        ]
        table = doc.add_table(rows=len(forces) + 1, cols=3)
        table.style = "Table Grid"
        _table_header(table, ["Competitive Force", "Intensity", "Key Factors"])
        for i, (label, key) in enumerate(forces, start=1):
            force_data = porter.get(key, {})
            row = table.rows[i].cells
            row[0].text = label
            rating = force_data.get("rating", "—")
            row[1].text = rating
            row[2].text = " | ".join(force_data.get("key_factors", [])[:3])
        doc.add_paragraph()

        if porter.get("strategic_implications"):
            _add_heading(doc, "Strategic Implications", level=3)
            for impl in porter["strategic_implications"]:
                _add_bullet(doc, impl)

    # 2.2 Competitor profiles
    profiles = state.get("competitor_profiles", [])
    if profiles:
        _add_heading(doc, "2.2 Market Leaders & Challengers", level=2)
        for profile in profiles:
            _add_heading(doc, profile["name"], level=3)
            _add_kv(doc, "Product", profile.get("product_summary", ""))
            _add_kv(doc, "Target Segment", profile.get("target_segment", ""))
            _add_kv(doc, "Pricing", profile.get("pricing_model", ""))
            if profile.get("funding"):
                _add_kv(doc, "Funding / Market Position", profile["funding"])

            if profile.get("core_features"):
                _add_body(doc, "Core Features:", bold=True)
                for f in profile["core_features"]:
                    _add_bullet(doc, f)

            if profile.get("strengths"):
                _add_body(doc, "Why Customers Choose Them:", bold=True)
                for s in profile["strengths"]:
                    _add_bullet(doc, s, color=GREEN)

            if profile.get("weaknesses"):
                _add_body(doc, "Customer Pain Points:", bold=True)
                for w in profile["weaknesses"]:
                    _add_bullet(doc, w, color=RED)

            if profile.get("differentiation_gaps"):
                _add_body(doc, "Gaps We Can Exploit:", bold=True)
                for gap in profile["differentiation_gaps"]:
                    _add_bullet(doc, gap, color=MEDIUM_BLUE)

            doc.add_paragraph()

    doc.add_page_break()


def _render_customer_insights(doc: Document, state: MICRAState) -> None:
    """Jobs-to-be-Done analysis — presented as Customer Insights, no framework label."""
    jtbd = _find_framework("jtbd", state.get("framework_outputs", []))
    if not jtbd:
        return

    _add_heading(doc, "3. Customer Insights", level=1)

    # Customer jobs
    jobs = jtbd.get("primary_jobs", [])
    if jobs:
        _add_heading(doc, "3.1 Jobs Customers Hire This Product to Do", level=2)
        _add_body(doc, "Customers don't buy products — they hire them to accomplish "
                  "specific goals. Understanding these jobs tells us exactly what to build.")
        doc.add_paragraph()

        for i, job in enumerate(jobs, 1):
            job_name = job.get("job_name", f"Job {i}")
            _add_heading(doc, f"Job {i}: {job_name}", level=3)

            when = job.get("when_context", "")
            want = job.get("desired_outcome", "")
            so_that = job.get("so_that", "")
            if when:
                _add_kv(doc, "When", when)
            if want:
                _add_kv(doc, "Goal", want)
            if so_that:
                _add_kv(doc, "So that", so_that)

            current = job.get("current_solutions", [])
            if current:
                _add_body(doc, "Current solutions used:", bold=True)
                for s in current:
                    _add_bullet(doc, s)

            pains = job.get("pain_points", [])
            if pains:
                _add_body(doc, "Pain points:", bold=True)
                for p in pains:
                    _add_bullet(doc, p, color=RED)

            metrics = job.get("success_metrics", [])
            if metrics:
                _add_body(doc, "Success looks like:", bold=True)
                for m in metrics:
                    _add_bullet(doc, m, color=GREEN)

            doc.add_paragraph()

    # Insight summary
    if jtbd.get("insight_summary"):
        _add_heading(doc, "Key Insight", level=2)
        p = doc.add_paragraph()
        r = p.add_run(jtbd["insight_summary"])
        r.font.color.rgb = DARK_BLUE
        r.bold = True

    # Personas
    _add_heading(doc, "3.2 Customer Segmentation by Job", level=2)
    for persona_key, persona_label in [("primary_persona", "Primary Persona"),
                                        ("secondary_persona", "Secondary Persona")]:
        persona = jtbd.get(persona_key, {})
        if persona:
            _add_heading(doc, persona_label, level=3)
            _add_kv(doc, "Role", persona.get("role", ""))
            if persona.get("age_range"):
                _add_kv(doc, "Profile", persona.get("age_range", ""))
            goals = persona.get("goals", "")
            if goals:
                _add_kv(doc, "Goals", goals if isinstance(goals, str)
                        else ", ".join(goals))
            challenges = persona.get("daily_challenges", "")
            if challenges:
                _add_kv(doc, "Challenges", challenges if isinstance(challenges, str)
                        else ", ".join(challenges))
            criteria = persona.get("decision_criteria", "")
            if criteria:
                _add_kv(doc, "Decision Criteria", criteria if isinstance(criteria, str)
                        else ", ".join(criteria))

    doc.add_page_break()


def _render_feature_strategy(doc: Document, state: MICRAState) -> None:
    """Feature prioritization from Kano — presented without naming the framework."""
    kano = _find_framework("kano", state.get("framework_outputs", []))
    if not kano:
        return

    _add_heading(doc, "4. Feature Strategy & Prioritization", level=1)
    _add_body(doc, "Features are classified by their impact on customer satisfaction, "
              "informing what must ship in the MVP vs. what can wait.")
    doc.add_paragraph()

    categories = [
        ("4.1 Table Stakes — Must Ship in MVP",
         "must_have",
         "Absence causes rejection. Presence is simply expected. "
         "These define the minimum bar to compete.",
         GREEN),
        ("4.2 Competitive Differentiators — More Is Better",
         "performance",
         "More of these = higher satisfaction. These are how we outcompete.",
         MEDIUM_BLUE),
        ("4.3 Innovation Opportunities — Unexpected Delight",
         "delighters",
         "Customers don't ask for these. Deliver one and create a wow moment.",
         ORANGE),
        ("4.4 What to Avoid",
         "indifferent",
         "These features neither help nor hurt — build them and you've wasted engineering time.",
         RED),
    ]

    for heading, key, description, color in categories:
        features = kano.get(key, [])
        if not features:
            continue
        _add_heading(doc, heading, level=2)
        _add_body(doc, description)
        for feat in features:
            if isinstance(feat, dict):
                name = feat.get("feature_name", "")
                desc = feat.get("description", "")
                text = f"{name}: {desc}" if desc else name
            else:
                text = str(feat)
            _add_bullet(doc, text, color=color)
        doc.add_paragraph()

    # Effort vs. Impact matrix (text representation)
    _add_heading(doc, "4.5 Feature Prioritization Matrix", level=2)
    _add_body(doc, "MVP focus: Quick Wins + 1-2 Major Projects")
    doc.add_paragraph()
    matrix_text = (
        "HIGH IMPACT\n"
        "  [Quick Wins]              [Major Projects]\n"
        "  Low effort, high impact   High effort, high impact\n"
        "  → Prioritize for MVP      → Select 1-2 for MVP\n"
        "\n"
        "  [Fill-Ins]                [Money Pits]\n"
        "  Low effort, low impact    High effort, low impact\n"
        "  → Include if convenient   → Avoid entirely\n"
        "LOW IMPACT"
    )
    p = doc.add_paragraph(matrix_text, style="No Spacing")
    p.runs[0].font.name = "Courier New"

    # Avoid / reverse features
    reverse = kano.get("reverse", [])
    if reverse:
        _add_heading(doc, "4.6 Features That Decrease Satisfaction — Avoid", level=2)
        for feat in reverse:
            if isinstance(feat, dict):
                _add_bullet(doc, feat.get("feature_name", str(feat)), color=RED)
            else:
                _add_bullet(doc, str(feat), color=RED)

    doc.add_page_break()


def _render_mvp_and_roadmap(doc: Document, state: MICRAState) -> None:
    _add_heading(doc, "5. MVP Definition & Roadmap", level=1)

    mvp = state.get("mvp_recommendation", {})
    roadmap = state.get("phased_roadmap", {})

    # 5.1 MVP scope
    _add_heading(doc, "5.1 MVP Scope", level=2)
    if mvp.get("differentiation_thesis"):
        p = doc.add_paragraph()
        r = p.add_run(f'Vision: "{mvp["differentiation_thesis"]}"')
        r.font.color.rgb = DARK_BLUE
        r.italic = True

    _add_kv(doc, "Target Customer", mvp.get("target_customer", "—"))
    _add_kv(doc, "North Star Metric", mvp.get("north_star_metric", "—"))
    _add_kv(doc, "Go-to-Market Entry", mvp.get("go_to_market_entry_point", "—"))

    if mvp.get("core_features"):
        _add_heading(doc, "MVP Must-Haves", level=3)
        for feat in mvp["core_features"]:
            _add_bullet(doc, feat, color=GREEN)

    if mvp.get("features_explicitly_excluded"):
        _add_heading(doc, "Explicitly NOT in v1 (scope control)", level=3)
        for feat in mvp["features_explicitly_excluded"]:
            _add_bullet(doc, feat, color=RED)

    # 5.2 Phased roadmap
    _add_heading(doc, "5.2 Post-MVP Roadmap", level=2)
    phases = [
        ("MVP", roadmap.get("mvp_phase", {})),
        ("Phase 2", roadmap.get("phase_2", {})),
        ("Phase 3", roadmap.get("phase_3", {})),
        ("Phase 4", roadmap.get("phase_4", {})),
    ]
    for phase_label, phase_data in phases:
        if not phase_data:
            continue
        timeline = phase_data.get("timeline", "")
        theme = phase_data.get("theme", phase_data.get("vision", ""))
        _add_heading(doc, f"{phase_label}: {timeline}", level=3)
        if theme:
            _add_body(doc, theme)

        # Focus areas or deliverables
        focus_areas = phase_data.get("focus_areas", [])
        deliverables = phase_data.get("key_deliverables", [])
        if focus_areas:
            for area in focus_areas:
                if isinstance(area, dict):
                    name = area.get("name", "")
                    effort = area.get("effort_percent", "")
                    label = f"{name} ({effort})" if effort else name
                    _add_bullet(doc, label)
                    for item in area.get("key_deliverables", [])[:3]:
                        p = doc.add_paragraph(f"    ◦ {item}", style="No Spacing")
                else:
                    _add_bullet(doc, str(area))
        elif deliverables:
            for d in deliverables:
                _add_bullet(doc, d)
        doc.add_paragraph()

    # 5.3 Go/No-Go checkpoints
    m3 = roadmap.get("go_no_go_month3", [])
    m6 = roadmap.get("go_no_go_month6", [])
    if m3 or m6:
        _add_heading(doc, "5.3 Go / No-Go Checkpoints", level=2)
        if m3:
            _add_heading(doc, "Month 3 Checkpoint", level=3)
            for c in m3:
                _add_bullet(doc, c)
        if m6:
            _add_heading(doc, "Month 6 Checkpoint (MVP)", level=3)
            for c in m6:
                _add_bullet(doc, c)

    # 5.4 6-month success criteria
    if mvp.get("success_criteria_6_months"):
        _add_heading(doc, "5.4 Success Criteria at 6 Months", level=2)
        for criterion in mvp["success_criteria_6_months"]:
            _add_bullet(doc, criterion)

    doc.add_page_break()


def _render_team_requirements(doc: Document, state: MICRAState) -> None:
    team = state.get("team_requirements", {})
    if not team:
        return

    _add_heading(doc, "6. Team Requirements & Capabilities", level=1)

    # 6.1 Core team
    core_roles = team.get("core_team_roles", [])
    if core_roles:
        _add_heading(doc, "6.1 Core Product Team (MVP Phase)", level=2)
        if team.get("total_core_headcount"):
            _add_kv(doc, "Total Headcount", team["total_core_headcount"])
        doc.add_paragraph()

        for role_data in core_roles:
            role = role_data.get("role", "")
            count = role_data.get("count", 1)
            focus = role_data.get("focus", "")
            skills = role_data.get("key_skills", [])

            _add_heading(doc, f"{role} × {count}", level=3)
            if focus:
                _add_body(doc, focus)
            if skills:
                _add_kv(doc, "Key Skills", " | ".join(skills))
            doc.add_paragraph()

    # 6.2 Critical skills
    critical = team.get("critical_skills", [])
    if critical:
        _add_heading(doc, "6.2 Critical Skills to Hire For", level=2)
        for skill in critical:
            _add_bullet(doc, skill, color=ORANGE)

    # 6.3 Hiring priorities
    priorities = team.get("hiring_priorities", [])
    if priorities:
        _add_heading(doc, "6.3 First 30-Day Hiring Priorities", level=2)
        for i, p in enumerate(priorities, 1):
            _add_bullet(doc, f"{i}. {p}")

    # 6.4 Extended team
    extended = team.get("extended_team_post_mvp", [])
    if extended:
        _add_heading(doc, "6.4 Extended Team (Post-MVP)", level=2)
        for role in extended:
            _add_bullet(doc, role)

    doc.add_page_break()


def _render_strategic_rationale(doc: Document, state: MICRAState) -> None:
    """Why enter this market — draws on SWOT without naming it."""
    swot = _find_framework("swot", state.get("framework_outputs", []))
    bbp = state.get("build_buy_partner_decision", {})
    tam = _find_framework("tam_sam_som", state.get("framework_outputs", []))

    _add_heading(doc, "7. Why Enter This Market?", level=1)

    # 7.1 Market opportunity
    _add_heading(doc, "7.1 Market Opportunity", level=2)
    if tam:
        for trend in tam.get("key_trends_driving_growth", [])[:4]:
            _add_bullet(doc, trend)

    # 7.2 Competitive gaps (from SWOT opportunities + BBP capability gaps)
    _add_heading(doc, "7.2 Gaps in the Current Market", level=2)
    if swot:
        for opp in swot.get("opportunities", []):
            _add_bullet(doc, opp, color=GREEN)
    if bbp.get("capability_gaps"):
        _add_heading(doc, "Capabilities we need to build or acquire:", level=3)
        for gap in bbp["capability_gaps"]:
            _add_bullet(doc, gap, color=ORANGE)

    # 7.3 Our strengths (from SWOT strengths)
    if swot and swot.get("strengths"):
        _add_heading(doc, "7.3 Our Strengths & Advantages", level=2)
        for strength in swot["strengths"]:
            _add_bullet(doc, strength, color=MEDIUM_BLUE)

    # 7.4 Threats to monitor
    if swot and swot.get("threats"):
        _add_heading(doc, "7.4 Threats to Monitor", level=2)
        for threat in swot["threats"]:
            _add_bullet(doc, threat, color=RED)

    # 7.5 Build/buy/partner cases
    if bbp.get("build_case"):
        _add_heading(doc, "7.5 Strategic Entry Options", level=2)
        _add_kv(doc, "Build In-House", bbp.get("build_case", ""))
        _add_kv(doc, "Acquire", bbp.get("buy_case", ""))
        _add_kv(doc, "Partner", bbp.get("partner_case", ""))
        if bbp.get("recommended_partners_or_targets"):
            _add_heading(doc, "Recommended Partners or Targets", level=3)
            for name in bbp["recommended_partners_or_targets"]:
                _add_bullet(doc, name)

    if swot and swot.get("recommended_strategic_posture"):
        doc.add_paragraph()
        _add_kv(doc, "Recommended Strategic Posture",
                swot["recommended_strategic_posture"])

    doc.add_page_break()


def _render_pricing_and_licensing(doc: Document, state: MICRAState) -> None:
    gtm = state.get("gtm_strategy", {})
    if not gtm:
        return

    _add_heading(doc, "8. Pricing & Licensing Models", level=1)

    pricing_options = gtm.get("pricing_options", [])
    if pricing_options:
        _add_heading(doc, "8.1 Pricing Options", level=2)
        for opt in pricing_options:
            model_name = opt.get("model_name", "")
            recommended = opt.get("recommended", False)
            label = f"{model_name}{'  ← Recommended' if recommended else ''}"
            _add_heading(doc, label, level=3)

            rationale = opt.get("rationale", "")
            if rationale:
                _add_body(doc, rationale)

            tiers = opt.get("tiers", [])
            if tiers:
                table = doc.add_table(rows=len(tiers) + 1, cols=3)
                table.style = "Table Grid"
                _table_header(table, ["Tier", "Price", "Includes"])
                for i, tier in enumerate(tiers, 1):
                    row = table.rows[i].cells
                    row[0].text = str(tier.get("name", ""))
                    row[1].text = str(tier.get("price", ""))
                    includes = tier.get("includes", "")
                    row[2].text = includes if isinstance(includes, str) else " | ".join(includes)
                doc.add_paragraph()

    # Revenue projections
    _add_heading(doc, "8.2 Revenue Projections", level=2)
    for label, key in [("Year 1 ARR", "revenue_year1"),
                        ("Year 3 ARR", "revenue_year3"),
                        ("Year 5 ARR", "revenue_year5")]:
        val = gtm.get(key, "")
        if val:
            _add_kv(doc, label, val)

    doc.add_page_break()


def _render_gtm_strategy(doc: Document, state: MICRAState) -> None:
    gtm = state.get("gtm_strategy", {})
    if not gtm:
        return

    _add_heading(doc, "9. Go-to-Market Strategy", level=1)

    # 9.1 Target segments
    segments = gtm.get("primary_target_segments", [])
    if segments:
        _add_heading(doc, "9.1 Target Customer Segments (Year 1)", level=2)
        for seg in segments:
            _add_bullet(doc, seg)

    # 9.2 Regional GTM
    regions = gtm.get("go_to_market_by_region", [])
    if regions:
        _add_heading(doc, "9.2 Regional Approach", level=2)
        for region_data in regions:
            region = region_data.get("region", "")
            _add_heading(doc, region, level=3)
            approach = region_data.get("go_to_market_approach", "")
            if approach:
                _add_body(doc, approach)
            channels = region_data.get("key_channels", [])
            if channels:
                _add_kv(doc, "Channels", " | ".join(channels))

    # 9.3 Sales & marketing
    _add_heading(doc, "9.3 Sales & Marketing", level=2)
    if gtm.get("sales_strategy"):
        _add_heading(doc, "Sales Strategy", level=3)
        _add_body(doc, gtm["sales_strategy"])
    channels = gtm.get("marketing_channels", [])
    if channels:
        _add_heading(doc, "Marketing Channels", level=3)
        for ch in channels:
            _add_bullet(doc, ch)

    # 9.4 Competitive messaging
    messaging = gtm.get("competitive_messaging", [])
    if messaging:
        _add_heading(doc, "9.4 Competitive Positioning", level=2)
        for item in messaging:
            competitor = item.get("competitor", "")
            messages = item.get("messages", [])
            if competitor and messages:
                _add_heading(doc, f"vs. {competitor}", level=3)
                for msg in messages:
                    _add_bullet(doc, f'"{msg}"', color=MEDIUM_BLUE)

    doc.add_page_break()


def _render_success_metrics(doc: Document, state: MICRAState) -> None:
    _add_heading(doc, "10. Success Metrics & KPIs", level=1)

    mvp = state.get("mvp_recommendation", {})
    gtm = state.get("gtm_strategy", {})

    # 10.1 North Star
    _add_heading(doc, "10.1 North Star Metric", level=2)
    if mvp.get("north_star_metric"):
        p = doc.add_paragraph()
        r = p.add_run(mvp["north_star_metric"])
        r.bold = True
        r.font.size = Pt(12)
        r.font.color.rgb = DARK_BLUE
        _add_body(doc, "This is the single metric that best captures value delivered to customers. "
                  "All feature decisions should ask: does this improve the north star?")

    # 10.2 Business metrics from revenue projections
    _add_heading(doc, "10.2 Business Metrics", level=2)
    for label, key in [("Year 1 Revenue Target", "revenue_year1"),
                        ("Year 3 Revenue Target", "revenue_year3")]:
        val = gtm.get(key, "")
        if val:
            _add_kv(doc, label, val)

    # 10.3 6-month success criteria
    criteria = mvp.get("success_criteria_6_months", [])
    if criteria:
        _add_heading(doc, "10.3 MVP Validation Milestones (6 Months)", level=2)
        for c in criteria:
            _add_bullet(doc, c, color=GREEN)

    doc.add_page_break()


def _render_risk_register(doc: Document, state: MICRAState) -> None:
    _add_heading(doc, "11. Risk Register", level=1)

    bbp = state.get("build_buy_partner_decision", {})
    swot = _find_framework("swot", state.get("framework_outputs", []))
    risks: list[str] = list(bbp.get("risk_factors", []))

    # Add SWOT threats as additional risks
    if swot:
        for threat in swot.get("threats", []):
            if threat not in risks:
                risks.append(threat)

    if risks:
        table = doc.add_table(rows=len(risks) + 1, cols=3)
        table.style = "Table Grid"
        _table_header(table, ["#", "Risk", "Source"])
        for i, risk in enumerate(risks, start=1):
            table.rows[i].cells[0].text = str(i)
            table.rows[i].cells[1].text = risk
            table.rows[i].cells[2].text = "Strategic" if i <= len(bbp.get("risk_factors", [])) else "Market"
    else:
        _add_body(doc, "Risk data was not captured in this run.")

    doc.add_page_break()


def _render_next_steps(doc: Document, state: MICRAState) -> None:
    roadmap = state.get("phased_roadmap", {})
    if not roadmap:
        return

    _add_heading(doc, "12. Next Steps & Action Items", level=1)

    _add_heading(doc, "12.1 Immediate Actions (Next 30 Days)", level=2)
    w12 = roadmap.get("week1_2_actions", [])
    w34 = roadmap.get("week3_4_actions", [])
    if w12:
        _add_heading(doc, "Weeks 1-2", level=3)
        for action in w12:
            _add_bullet(doc, f"☐  {action}")
    if w34:
        _add_heading(doc, "Weeks 3-4", level=3)
        for action in w34:
            _add_bullet(doc, f"☐  {action}")

    _add_heading(doc, "12.2 90-Day Milestones", level=2)
    m1 = roadmap.get("month1_milestones", [])
    m3 = roadmap.get("month3_milestones", [])
    if m1:
        _add_heading(doc, "Month 1", level=3)
        for m in m1:
            _add_bullet(doc, m)
    if m3:
        _add_heading(doc, "Month 3", level=3)
        for m in m3:
            _add_bullet(doc, m)


def _render_sources(doc: Document, cited_sources: list[dict]) -> None:
    doc.add_page_break()
    _add_heading(doc, "13. Sources & Citations", level=1)
    if cited_sources:
        for i, source in enumerate(cited_sources, start=1):
            p = doc.add_paragraph()
            p.add_run(f"[{i}] ").bold = True
            p.add_run(f"{source['title']} ")
            run = p.add_run(f"({source['type']})")
            run.font.color.rgb = MEDIUM_BLUE
            p.add_run(f"\n    {source['url']}")
    else:
        _add_body(doc, "Source citation data was not available for this run.")


# ── Main node ──────────────────────────────────────────────────────────────

def report_generator_node(state: MICRAState) -> dict:
    """
    Generate the Word report from accumulated state.

    LEARNING: This node only READS state — it writes nothing back except
    the report file path. It's a pure consumer of all previous nodes' work.

    The report structure mirrors a professional market intelligence document:
    Executive summary → Market → Competitive landscape → Customer insights →
    Feature strategy → MVP → Team → Rationale → Pricing → GTM →
    Metrics → Risks → Next steps → Sources

    Framework names are intentionally omitted from section headings.
    A good analyst presents insights, not methodology labels.
    """
    console.print("\n[bold cyan]Phase 4: Generating Report[/bold cyan]")

    plan = state.get("research_plan", {})
    market = plan.get("target_market", "market")
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    market_safe = re.sub(r"[^\w\s]", "_", market)
    market_safe = re.sub(r"\s+", "_", market_safe).lower().strip("_")
    filename = f"micra_{market_safe}_{timestamp}.docx"
    output_path = os.path.join("outputs", filename)
    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    console.print("  Building sections...", end=" ")

    _render_title_page(doc, state)
    _render_executive_summary(doc, state)
    _render_market_overview(doc, state)
    _render_competitive_landscape(doc, state)
    _render_customer_insights(doc, state)
    _render_feature_strategy(doc, state)
    _render_mvp_and_roadmap(doc, state)
    _render_team_requirements(doc, state)
    _render_strategic_rationale(doc, state)
    _render_pricing_and_licensing(doc, state)
    _render_gtm_strategy(doc, state)
    _render_success_metrics(doc, state)
    _render_risk_register(doc, state)
    _render_next_steps(doc, state)

    cited_sources = _collect_all_sources(state)
    _render_sources(doc, cited_sources)

    doc.save(output_path)
    console.print("[green]✓[/green]")

    section_count = sum([
        bool(state.get("build_buy_partner_decision")),
        bool(state.get("mvp_recommendation")),
        bool(state.get("gtm_strategy")),
        bool(state.get("team_requirements")),
        bool(state.get("phased_roadmap")),
        bool(state.get("framework_outputs")),
        bool(state.get("competitor_profiles")),
    ])
    console.print(
        f"[green]✓[/green] Report saved: [bold]{output_path}[/bold] "
        f"({section_count} data sections, {len(cited_sources)} sources cited)"
    )

    return {
        "report_path": output_path,
        "messages": [
            f"[report_generator] Report generated: {output_path} "
            f"({len(cited_sources)} sources cited, {section_count} sections)"
        ],
    }
